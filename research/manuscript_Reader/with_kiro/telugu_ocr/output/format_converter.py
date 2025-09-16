"""Format conversion utilities for different output formats."""

import csv
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
from docx.shared import Inches

from ..models.data_models import OCRResult
from ..utils.exceptions import OutputGenerationException
from ..utils.logging_config import get_logger_config


class FormatConverter:
    """Converts OCR results to various output formats."""
    
    def __init__(self):
        """Initialize format converter."""
        self.logger_config = get_logger_config()
        self.process_logger = self.logger_config.get_process_logger()
        self.error_logger = self.logger_config.get_error_logger()
    
    def text_to_csv_rows(self, text: str, preserve_structure: bool = True) -> List[List[str]]:
        """
        Convert text to CSV-compatible rows.
        
        Args:
            text: Input text
            preserve_structure: Whether to preserve line structure
            
        Returns:
            List of CSV rows
        """
        if not text or not text.strip():
            return [['[No text detected]']]
        
        try:
            if preserve_structure:
                # Split by lines and preserve structure
                lines = text.split('\n')
                rows = []
                
                for line in lines:
                    line = line.strip()
                    if line:
                        # Try to detect table-like structure
                        if self._looks_like_table_row(line):
                            # Split by common delimiters
                            cells = self._split_table_row(line)
                            rows.append(cells)
                        else:
                            # Regular text line
                            rows.append([line])
                    else:
                        # Empty line - add as separator
                        rows.append([''])
                
                return rows if rows else [['[No text detected]']]
            else:
                # Simple line-by-line conversion
                lines = [line.strip() for line in text.split('\n') if line.strip()]
                return [[line] for line in lines] if lines else [['[No text detected]']]
                
        except Exception as e:
            self.error_logger.error(f"Error converting text to CSV rows: {str(e)}")
            return [[text]]  # Fallback to single cell
    
    def _looks_like_table_row(self, line: str) -> bool:
        """Check if a line looks like a table row."""
        # Common table indicators
        table_indicators = ['|', '\t', '  ', ':', '-']
        
        # Count potential separators
        separator_count = 0
        for indicator in table_indicators:
            if indicator in line:
                separator_count += line.count(indicator)
        
        # If line has multiple separators, likely a table
        return separator_count >= 2
    
    def _split_table_row(self, line: str) -> List[str]:
        """Split a table-like row into cells."""
        # Try different delimiters in order of preference
        delimiters = ['|', '\t', '  ', ':']
        
        for delimiter in delimiters:
            if delimiter in line:
                cells = [cell.strip() for cell in line.split(delimiter)]
                # Filter out empty cells from splitting
                cells = [cell for cell in cells if cell]
                if len(cells) > 1:
                    return cells
        
        # Fallback: split by multiple spaces
        import re
        cells = re.split(r'\s{2,}', line.strip())
        return [cell.strip() for cell in cells if cell.strip()]
    
    def create_structured_csv(self, results: List[OCRResult], output_path: str) -> str:
        """
        Create a structured CSV with all pages and metadata.
        
        Args:
            results: List of OCR results
            output_path: Output file path
            
        Returns:
            Path to created CSV file
        """
        try:
            with open(output_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                
                # Write header
                writer.writerow(['Page', 'Line', 'Text', 'Confidence', 'Processing_Time'])
                
                for result in results:
                    lines = result.text.split('\n')
                    for line_num, line in enumerate(lines, 1):
                        if line.strip():
                            writer.writerow([
                                result.page_number,
                                line_num,
                                line.strip(),
                                f"{result.confidence_score:.3f}",
                                f"{result.processing_time:.3f}"
                            ])
            
            self.process_logger.info(f"Created structured CSV: {output_path}")
            return output_path
            
        except Exception as e:
            raise OutputGenerationException(f"Failed to create structured CSV: {str(e)}")
    
    def create_excel_workbook(self, all_results: Dict[str, List[OCRResult]], output_path: str) -> str:
        """
        Create Excel workbook with multiple sheets for different engines.
        
        Args:
            all_results: Dictionary mapping engine names to results
            output_path: Output file path
            
        Returns:
            Path to created Excel file
        """
        try:
            with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                # Create summary sheet
                self._create_summary_sheet(all_results, writer)
                
                # Create sheet for each engine
                for engine_name, results in all_results.items():
                    self._create_engine_sheet(engine_name, results, writer)
                
                # Create comparison sheet
                self._create_comparison_sheet(all_results, writer)
            
            self.process_logger.info(f"Created Excel workbook: {output_path}")
            return output_path
            
        except Exception as e:
            raise OutputGenerationException(f"Failed to create Excel workbook: {str(e)}")
    
    def _create_summary_sheet(self, all_results: Dict[str, List[OCRResult]], writer):
        """Create summary sheet in Excel workbook."""
        summary_data = []
        
        for engine_name, results in all_results.items():
            if results:
                avg_confidence = sum(r.confidence_score for r in results) / len(results)
                total_time = sum(r.processing_time for r in results)
                total_chars = sum(len(r.text) for r in results)
                
                summary_data.append({
                    'Engine': engine_name,
                    'Pages': len(results),
                    'Avg_Confidence': f"{avg_confidence:.3f}",
                    'Total_Time': f"{total_time:.3f}",
                    'Total_Characters': total_chars,
                    'Chars_Per_Second': f"{total_chars/total_time:.1f}" if total_time > 0 else "N/A"
                })
        
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
    
    def _create_engine_sheet(self, engine_name: str, results: List[OCRResult], writer):
        """Create sheet for specific engine results."""
        engine_data = []
        
        for result in results:
            lines = result.text.split('\n')
            for line_num, line in enumerate(lines, 1):
                if line.strip():
                    engine_data.append({
                        'Page': result.page_number,
                        'Line': line_num,
                        'Text': line.strip(),
                        'Confidence': f"{result.confidence_score:.3f}",
                        'Processing_Time': f"{result.processing_time:.3f}"
                    })
        
        if engine_data:
            engine_df = pd.DataFrame(engine_data)
            # Limit sheet name length
            sheet_name = engine_name[:31] if len(engine_name) > 31 else engine_name
            engine_df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    def _create_comparison_sheet(self, all_results: Dict[str, List[OCRResult]], writer):
        """Create comparison sheet showing results side by side."""
        if not all_results:
            return
        
        # Get all page numbers
        all_pages = set()
        for results in all_results.values():
            all_pages.update(r.page_number for r in results)
        
        comparison_data = []
        
        for page_num in sorted(all_pages):
            row = {'Page': page_num}
            
            for engine_name, results in all_results.items():
                # Find result for this page
                page_result = next((r for r in results if r.page_number == page_num), None)
                
                if page_result:
                    # Truncate text for comparison view
                    text_preview = page_result.text[:100] + "..." if len(page_result.text) > 100 else page_result.text
                    text_preview = text_preview.replace('\n', ' ')
                    row[f"{engine_name}_Text"] = text_preview
                    row[f"{engine_name}_Confidence"] = f"{page_result.confidence_score:.3f}"
                else:
                    row[f"{engine_name}_Text"] = "[No result]"
                    row[f"{engine_name}_Confidence"] = "0.000"
            
            comparison_data.append(row)
        
        if comparison_data:
            comparison_df = pd.DataFrame(comparison_data)
            comparison_df.to_excel(writer, sheet_name='Comparison', index=False)
    
    def create_word_document(self, results: List[OCRResult], output_path: str, 
                           engine_name: str = "OCR") -> str:
        """
        Create Word document with OCR results.
        
        Args:
            results: List of OCR results
            output_path: Output file path
            engine_name: Name of OCR engine
            
        Returns:
            Path to created Word document
        """
        try:
            doc = Document()
            
            # Add title
            title = doc.add_heading(f'OCR Results - {engine_name}', 0)
            
            # Add metadata
            doc.add_paragraph(f'Generated: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Pages: {len(results)}')
            doc.add_paragraph('')
            
            # Add results for each page
            for result in results:
                # Page header
                page_heading = doc.add_heading(f'Page {result.page_number}', level=1)
                
                # Metadata
                metadata_para = doc.add_paragraph()
                metadata_para.add_run(f'Confidence: {result.confidence_score:.3f} | ')
                metadata_para.add_run(f'Processing Time: {result.processing_time:.3f}s')
                
                # Text content
                if result.text.strip():
                    doc.add_paragraph(result.text)
                else:
                    doc.add_paragraph('[No text detected]')
                
                # Add page break except for last page
                if result != results[-1]:
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            
            self.process_logger.info(f"Created Word document: {output_path}")
            return output_path
            
        except Exception as e:
            raise OutputGenerationException(f"Failed to create Word document: {str(e)}")
    
    def create_json_export(self, all_results: Dict[str, List[OCRResult]], 
                          output_path: str, include_metadata: bool = True) -> str:
        """
        Create comprehensive JSON export of all results.
        
        Args:
            all_results: Dictionary mapping engine names to results
            output_path: Output file path
            include_metadata: Whether to include detailed metadata
            
        Returns:
            Path to created JSON file
        """
        try:
            export_data = {
                'export_info': {
                    'generated_time': pd.Timestamp.now().isoformat(),
                    'total_engines': len(all_results),
                    'total_pages': max(len(results) for results in all_results.values()) if all_results else 0
                },
                'engines': {}
            }
            
            for engine_name, results in all_results.items():
                engine_data = {
                    'engine_name': engine_name,
                    'total_pages': len(results),
                    'pages': []
                }
                
                if include_metadata:
                    # Add engine-level statistics
                    if results:
                        confidences = [r.confidence_score for r in results]
                        times = [r.processing_time for r in results]
                        
                        engine_data['statistics'] = {
                            'avg_confidence': sum(confidences) / len(confidences),
                            'min_confidence': min(confidences),
                            'max_confidence': max(confidences),
                            'total_processing_time': sum(times),
                            'avg_processing_time': sum(times) / len(times)
                        }
                
                # Add page results
                for result in results:
                    page_data = {
                        'page_number': result.page_number,
                        'text': result.text,
                        'confidence_score': result.confidence_score,
                        'processing_time': result.processing_time
                    }
                    
                    if include_metadata:
                        page_data['metadata'] = {
                            'timestamp': result.timestamp.isoformat() if result.timestamp else None,
                            'text_length': len(result.text),
                            'line_count': len(result.text.split('\n')),
                            'bounding_boxes_count': len(result.bounding_boxes)
                        }
                        
                        # Add bounding boxes if available
                        if result.bounding_boxes:
                            page_data['bounding_boxes'] = [
                                {
                                    'x': bbox.x,
                                    'y': bbox.y,
                                    'width': bbox.width,
                                    'height': bbox.height
                                }
                                for bbox in result.bounding_boxes
                            ]
                    
                    engine_data['pages'].append(page_data)
                
                export_data['engines'][engine_name] = engine_data
            
            # Save JSON
            with open(output_path, 'w', encoding='utf-8') as jsonfile:
                json.dump(export_data, jsonfile, indent=2, ensure_ascii=False)
            
            self.process_logger.info(f"Created JSON export: {output_path}")
            return output_path
            
        except Exception as e:
            raise OutputGenerationException(f"Failed to create JSON export: {str(e)}")